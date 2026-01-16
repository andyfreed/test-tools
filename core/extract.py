import io
import re
import zipfile
from typing import Any, Dict, List, Tuple

from docx import Document

from .utils import ENCODING_WARNING, normalize_text


REPLACEMENT_CHAR = "\ufffd"

ANSWER_KEY_ANNOTATION_PATTERN = re.compile(r"\s*\[[^\]]+\]\s*$")
QUESTION_START_PATTERNS = [
    re.compile(r"^\s*(\d{1,4})[.)]\s+\S"),
    re.compile(r"^\s*Q(\d{1,4})[:.)]\s+\S", flags=re.IGNORECASE),
    # Page-reference markers like [p1-4], [p2-6]
    re.compile(r"^\s*\[[^\]]*p\d+[^\]]*\]", flags=re.IGNORECASE),
]
OPTION_PATTERNS = [
    re.compile(r"^\s*[A-D][.)]\s+\S", flags=re.IGNORECASE),
    re.compile(r"^\s*\(?[A-D]\)\s+\S", flags=re.IGNORECASE),
]
ANSWER_KEY_PATTERN = re.compile(r"^\s*(?:Q\s*)?(\d{1,4})\s*[.)]?\s*[:=\-]?\s*([A-D])\b", flags=re.IGNORECASE)


def _add_warning(warnings: List[str], message: str) -> None:
    if message not in warnings:
        warnings.append(message)


def _strip_answer_key_annotation(text: str) -> str:
    """Remove trailing bracketed annotations like [Chp 1] from answer key lines."""
    return re.sub(ANSWER_KEY_ANNOTATION_PATTERN, "", text or "")


def _decode_text_bytes(content: bytes) -> Tuple[str, bool]:
    """
    Decode bytes using a prioritized encoding list, avoiding heavy use of U+FFFD.

    Returns decoded text and whether encoding artifacts were detected.
    """
    encodings = ["utf-8-sig", "utf-8", "cp1252", "latin-1"]
    for enc in encodings:
        decoded = content.decode(enc, errors="replace")
        replacement_count = decoded.count(REPLACEMENT_CHAR)
        replacement_ratio = replacement_count / max(len(decoded), 1)
        if replacement_count == 0 or replacement_ratio <= 0.01:
            artifacts = enc not in ("utf-8-sig", "utf-8") or replacement_count > 0
            return decoded, artifacts

    # Fallback: use the last encoding even if replacements remain.
    decoded = content.decode(encodings[-1], errors="replace")
    return decoded, True


def _analyze_question_patterns(texts: List[str]) -> Dict[str, int]:
    """Compute debug counts for question starts, options, and answer key entries."""
    question_starts = 0
    option_lines = 0
    answer_key_entries = 0

    for raw in texts:
        if any(pat.match(raw) for pat in QUESTION_START_PATTERNS):
            question_starts += 1
        if any(pat.match(raw) for pat in OPTION_PATTERNS):
            option_lines += 1
        cleaned = _strip_answer_key_annotation(raw)
        if ANSWER_KEY_PATTERN.match(cleaned):
            answer_key_entries += 1

    return {
        "total_lines": len(texts),
        "question_starts": question_starts,
        "option_lines": option_lines,
        "answer_key_entries": answer_key_entries,
    }


def _detect_tracked_changes(content: bytes) -> bool:
    """Inspect word/document.xml for tracked change tags with visible text."""
    try:
        with zipfile.ZipFile(io.BytesIO(content)) as zf:
            xml_bytes = zf.read("word/document.xml")
            xml_text = xml_bytes.decode("utf-8", errors="ignore")
            pattern = re.compile(
                r"<w:(ins|del|moveFrom|moveTo)\b[^>]*>.*?<w:t\b[^>]*>\s*\S",
                flags=re.IGNORECASE | re.DOTALL,
            )
            return bool(pattern.search(xml_text))
    except Exception:
        return False


def extract_docx(content: bytes, filename: str) -> Dict[str, Any]:
    """Extract paragraphs, highlight flags, and tracked changes flag from a DOCX file."""
    has_tracked_changes = _detect_tracked_changes(content)
    doc = Document(io.BytesIO(content))
    paragraphs: List[Dict[str, Any]] = []
    line_index = 0

    def append_paragraph(text: str, has_highlight: bool = False) -> None:
        nonlocal line_index
        if not text:
            return
        paragraphs.append(
            {
                "i": line_index,
                "line_index": line_index,
                "text": text,
                "has_highlight": bool(has_highlight),
            }
        )
        line_index += 1

    for paragraph in doc.paragraphs:
        text = normalize_text(paragraph.text or "")
        has_highlight = any(
            run.text and run.text.strip() and run.font.highlight_color is not None
            for run in paragraph.runs
        )
        append_paragraph(text, has_highlight)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text = normalize_text(para.text or "")
                    # Table paragraphs rarely have highlight; treat as False by default.
                    append_paragraph(text, has_highlight=False)

    debug_counts = _analyze_question_patterns([p["text"] for p in paragraphs])
    return {
        "source_filename": filename,
        "content_type": "docx",
        "paragraphs": paragraphs,
        "has_tracked_changes": has_tracked_changes,
        "debug_counts": debug_counts,
    }


def extract_txt(content: bytes, filename: str) -> Dict[str, Any]:
    """Extract non-empty lines from a text file."""
    decoded, had_artifacts = _decode_text_bytes(content)
    warnings: List[str] = []
    lines: List[Dict[str, Any]] = []
    line_index = 0
    for raw in decoded.splitlines():
        text = normalize_text(raw, warnings=warnings)
        if not text:
            continue
        lines.append({"i": line_index, "line_index": line_index, "text": text})
        line_index += 1
    if had_artifacts:
        _add_warning(warnings, ENCODING_WARNING)
    debug_counts = _analyze_question_patterns([l["text"] for l in lines])
    return {
        "source_filename": filename,
        "content_type": "txt",
        "lines": lines,
        "warnings": warnings,
        "debug_counts": debug_counts,
    }


def build_document_signals(files: List[Any]) -> List[Dict[str, Any]]:
    """Convert uploaded files into document signal structures."""
    signals: List[Dict[str, Any]] = []
    for uploaded in files:
        filename = getattr(uploaded, "name", "uploaded_file")
        content = uploaded.read() if hasattr(uploaded, "read") else uploaded
        if isinstance(content, str):
            content = content.encode("utf-8")
        lower = filename.lower()
        if lower.endswith(".docx"):
            signals.append(extract_docx(content, filename))
        elif lower.endswith(".txt"):
            signals.append(extract_txt(content, filename))
        else:
            raise ValueError(f"Unsupported file type for {filename}. Use .docx or .txt.")
    return signals
