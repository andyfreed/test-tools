import io
import os
import tempfile

from docx import Document
from docx.enum.text import WD_COLOR_INDEX

from core.export_csv import build_csv_bytes
from core.extract import extract_docx, extract_txt
from core.utils import normalize_text
from core.validate import validate_parsed_questions


def _build_sample_docx() -> bytes:
    doc = Document()
    doc.add_paragraph("Sample Question 1")
    doc.add_paragraph("*Option A")
    p = doc.add_paragraph("Option B")
    run = p.add_run(" (highlighted)")
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    doc.add_paragraph("Option C")
    doc.add_paragraph("Option D")
    file_obj = io.BytesIO()
    doc.save(file_obj)
    return file_obj.getvalue()


def test_extract_docx_highlight_detection():
    content = _build_sample_docx()
    signal = extract_docx(content, "sample.docx")
    assert signal["content_type"] == "docx"
    has_highlight = any(p["has_highlight"] for p in signal["paragraphs"])
    assert has_highlight, "Expected at least one highlighted paragraph"


def test_extract_txt_lines():
    content = b"Q1: Sample?\nA) One\n\nB) Two\n"
    signal = extract_txt(content, "sample.txt")
    texts = [l["text"] for l in signal["lines"]]
    assert "Q1: Sample?" in texts
    assert len(texts) == 3  # skips empty line


def test_validation():
    good = {
        "category": "Demo",
        "questions": [
            {
                "number": 1,
                "title": "What is 2+2?",
                "options": ["1", "2", "3", "4"],
                "correct_index": 3,
                "detected_answer_method": "asterisk",
                "confidence": "high",
                "warnings": [],
                "source_refs": [{"kind": "paragraph", "index": 0}],
            }
        ],
    }
    assert validate_parsed_questions(good) == []

    bad = {"category": "Demo", "questions": [{"number": 0, "title": "", "options": [], "correct_index": 5}]}
    errors = validate_parsed_questions(bad)
    assert errors, "Validation should find errors for bad payload"


def test_normalize_text_mojibake_cleanup():
    assert normalize_text("Reduce the modelâ€™s complexity") == "Reduce the model's complexity"
    assert normalize_text("Reduce the modelâs complexity") == "Reduce the model's complexity"
    assert normalize_text("Precision Ã— Recall") == "Precision × Recall"


def test_asterisk_detection_in_extraction():
    content = b"1. What is X?\nA) Option one\n*B) Correct option*\nC) Option three\nD) Option four\n"
    signal = extract_txt(content, "sample.txt")
    asterisk_lines = [l for l in signal["lines"] if l.get("has_asterisk")]
    assert asterisk_lines, "Expected at least one line flagged with has_asterisk"
    assert "Correct option" in asterisk_lines[0]["text"]


def test_validation_catches_missing_confidence():
    data = {
        "category": "Demo",
        "questions": [
            {
                "number": 1,
                "title": "Test?",
                "options": ["A", "B", "C", "D"],
                "correct_index": 0,
                "detected_answer_method": "inferred",
                "warnings": [],
                "source_refs": [],
            }
        ],
    }
    errors = validate_parsed_questions(data)
    assert any("confidence" in e for e in errors), "Should flag missing confidence"


def test_validation_catches_duplicate_options():
    data = {
        "category": "Demo",
        "questions": [
            {
                "number": 1,
                "title": "Test?",
                "options": ["Same", "Same", "Different", "Other"],
                "correct_index": 0,
                "detected_answer_method": "inferred",
                "confidence": "low",
                "warnings": [],
                "source_refs": [],
            }
        ],
    }
    errors = validate_parsed_questions(data)
    assert any("duplicate" in e.lower() for e in errors), "Should detect duplicate options"


def test_docx_table_highlight_detection():
    """Highlights in table cells should be detected, not ignored."""
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    p = cell.paragraphs[0]
    run = p.add_run("Highlighted in table")
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    file_obj = io.BytesIO()
    doc.save(file_obj)
    signal = extract_docx(file_obj.getvalue(), "table.docx")
    highlighted = [p for p in signal["paragraphs"] if p.get("has_highlight")]
    assert highlighted, "Should detect highlights in table cells"


def test_csv_export_uses_utf8_bom():
    payload = {
        "questions": [
            {
                "number": 1,
                "title": "F1 score’s formula",
                "options": [
                    "(Precision + Recall) ÷ 2",
                    "2 × (Precision × Recall) ÷ (Precision + Recall)",
                    "Precision × Recall",
                    "Precision + Recall",
                ],
                "correct_index": 1,
                "detected_answer_method": "inferred",
                "warnings": [],
                "source_refs": [{"kind": "line", "index": 0}],
            }
        ]
    }
    data = build_csv_bytes(payload, "test")
    assert data.startswith(b"\xef\xbb\xbf")
    decoded = data.decode("utf-8-sig")
    assert "\u00d7" in decoded
    assert "\u00f7" in decoded
